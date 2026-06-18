from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "final_report" / "RiskRadar_AI_Final_Report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 35, 55)
MUTED = RGBColor(90, 96, 105)
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, after=6, before=0, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run_font(r, size=11, bold=bold, color=color or INK)
    return p


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)


def add_numbered(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for run in p.runs:
        set_run_font(run, size=size, bold=True, color=color)
    return p


def add_callout(doc, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_paragraph(doc, "", after=4)


def add_key_value_table(doc, rows: list[tuple[str, str]], widths=(2500, 6860)) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, list(widths))
    for i, (label, value) in enumerate(rows):
        label_cell = table.cell(i, 0)
        value_cell = table.cell(i, 1)
        set_cell_fill(label_cell, LIGHT_GRAY)
        for cell, text, bold in ((label_cell, label, True), (value_cell, value, False)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=10.5, bold=bold, color=INK)
    add_paragraph(doc, "", after=4)


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths_dxa)
    hdr = table.rows[0]
    hdr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_fill(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10, bold=True, color=INK)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.8, color=INK)
    set_table_width(table, widths_dxa)
    add_paragraph(doc, "", after=4)


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9.5, italic=True, color=MUTED)


def add_image(doc, rel_path: str, caption: str, width_in: float = 6.25) -> None:
    path = ROOT / rel_path
    if not path.exists():
        add_paragraph(doc, f"Screenshot missing: {rel_path}", color=RGBColor(155, 28, 28))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def add_code_block(doc, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Courier New", size=9.5, color=INK)
    add_paragraph(doc, "", after=4)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("RiskRadar AI Final Report")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    add_page_number(fp)


def add_cover(doc: Document) -> None:
    add_paragraph(doc, "", after=18)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("FINAL PROJECT REPORT")
    set_run_font(r, size=12, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("RiskRadar AI")
    set_run_font(r, size=28, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("AI-Powered Fraud Transaction Investigation Assistant")
    set_run_font(r, size=15, color=DARK_BLUE)

    add_key_value_table(
        doc,
        [
            ("Capstone", "IIT Roorkee AIOps Capstone"),
            ("Theme", "Theme 13 - AI Fraud Transaction Investigation Assistant"),
            ("Prepared", "June 2026"),
            ("Internal readiness target", "25 June 2026"),
            ("Updated submission deadline", "30 June 2026"),
            ("Repository", "https://github.com/Alkarth2010/RiskRadar_AI"),
        ],
        widths=(2600, 6760),
    )
    add_callout(
        doc,
        "Report purpose",
        "This report summarizes the design, implementation, testing, deployment, and submission readiness of RiskRadar AI, an agentic fraud investigation assistant built with Streamlit, LangGraph, LangChain, FAISS, and AWS.",
    )
    doc.add_page_break()


def build() -> None:
    doc = Document()
    setup_styles(doc)
    setup_header_footer(doc)
    add_cover(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_paragraph(
        doc,
        "RiskRadar AI is an AI-powered fraud transaction investigation assistant designed to help fraud analysts investigate suspicious bank transactions faster, more consistently, and with clearer auditability. The system generates alerts from synthetic transaction data, runs a structured multi-agent investigation workflow, retrieves policy-grounded evidence, recommends analyst actions, and records human decisions for future evaluation.",
    )
    add_paragraph(
        doc,
        "The final solution includes a Streamlit analyst dashboard, deterministic risk scoring, LangGraph orchestration, FAISS-backed retrieval over fraud policy documents, optional Gemini-assisted summaries, AWS EC2 deployment, S3-backed file storage, and optional Docker containerization.",
    )
    add_callout(
        doc,
        "Current readiness",
        "Core implementation, dashboard, RAG, workflow, human approval, AWS deployment, S3 storage, Docker support, and screenshot evidence are complete. Remaining work is limited to final presentation, demo recording, and submission packaging.",
    )

    add_heading(doc, "2. Problem Statement", 1)
    add_paragraph(
        doc,
        "Fraud analysts often review high volumes of suspicious transactions under time pressure. Manual investigations can become inconsistent when analysts must separately inspect transaction attributes, triggered rules, policy documents, evidence sources, and final action recommendations.",
    )
    add_paragraph(
        doc,
        "RiskRadar AI addresses this gap by creating a repeatable investigation workflow that combines deterministic fraud scoring, policy evidence retrieval, explainable reasoning, and human-in-the-loop approval.",
    )

    add_heading(doc, "3. Objectives", 1)
    for item in [
        "Generate a prioritized alert queue from synthetic transaction data.",
        "Investigate selected alerts through a structured agentic workflow.",
        "Ground investigation reasoning in fraud policy documents.",
        "Produce explainable recommendations: ESCALATE, MONITOR, or APPROVE.",
        "Allow analysts to accept, override, or annotate system recommendations.",
        "Persist feedback for audit, review, and future evaluation.",
        "Deploy the dashboard on AWS and support optional Docker-based reproducibility.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Solution Overview", 1)
    add_paragraph(
        doc,
        "The system is organized as an end-to-end analyst workflow. Alerts are generated from transaction data and shown in Streamlit. Once an analyst selects an alert, a LangGraph workflow fans out into specialist investigation branches: risk scoring, policy evidence retrieval, and behavioral pattern review. The evidence is fused into a structured investigation result and a recommendation. The analyst then records the final decision and optional note.",
    )
    add_code_block(
        doc,
        [
            "Synthetic transactions",
            "  -> generated alert queue",
            "  -> selected alert",
            "  -> parallel LangGraph investigation",
            "  -> risk reasoning + policy evidence + recommendation",
            "  -> analyst decision",
            "  -> feedback log and decision history",
        ],
    )
    add_image(
        doc,
        "docs/screenshots/Demo/01_alert_queue_initial.jpg",
        "Figure 1. Streamlit alert queue with priority metrics and alert selection.",
    )

    add_heading(doc, "5. System Architecture", 1)
    add_table(
        doc,
        ["Layer", "Purpose", "Key Technology"],
        [
            ["Data layer", "Synthetic transactions, fraud policies, and analyst feedback.", "CSV files, policy text files, S3"],
            ["Alert layer", "Generate alerts and queue priority from transaction risk indicators.", "Pandas, Pydantic"],
            ["Workflow layer", "Coordinate investigation branches and merge evidence.", "LangGraph"],
            ["RAG layer", "Retrieve supporting fraud policy context and source documents.", "LangChain, FAISS, sentence-transformers"],
            ["UI layer", "Provide analyst queue, investigation result, and human decision capture.", "Streamlit"],
            ["Deployment layer", "Run the app in cloud and local reproducible environments.", "AWS EC2, S3, Docker"],
        ],
        [1900, 5060, 2400],
    )
    add_heading(doc, "5.1 LangGraph Workflow", 2)
    add_code_block(
        doc,
        [
            "alert_intake",
            "  -> risk_scoring",
            "  -> policy_evidence",
            "  -> behavioral_pattern",
            "  -> evidence_fusion",
            "  -> recommendation",
            "  -> human decision pending",
        ],
    )
    add_paragraph(
        doc,
        "The workflow uses a fan-out/fan-in pattern. After alert intake, risk scoring, policy evidence retrieval, and behavioral analysis run as separate specialist branches. Their outputs merge at evidence fusion before the final recommendation is generated.",
    )

    add_heading(doc, "6. Fraud Risk and Recommendation Logic", 1)
    add_table(
        doc,
        ["Policy", "Trigger", "Weight"],
        [
            ["High Value Transaction", "Transaction amount greater than INR 100,000", "30"],
            ["Geographic Anomaly", "High-risk country or unusual location signal", "25"],
            ["Velocity and Burst Detection", "More than 10 transactions in 10 minutes", "25"],
            ["Device and Payment Instrument", "New device or payment instrument indicator", "20"],
        ],
        [3000, 4960, 1400],
    )
    add_table(
        doc,
        ["Risk Level", "Condition", "System Recommendation"],
        [
            ["HIGH", "Score >= 70 or multiple strong fraud indicators", "ESCALATE"],
            ["MEDIUM", "Score >= 20 with at least one meaningful indicator", "MONITOR"],
            ["LOW", "No significant triggered policy risk", "APPROVE"],
        ],
        [1800, 5260, 2300],
    )
    add_image(
        doc,
        "docs/screenshots/Demo/05_risk_reasoning_and_policies.jpg",
        "Figure 2. Investigation result showing risk reasoning and triggered policies.",
    )

    add_heading(doc, "7. Policy-Grounded RAG Pipeline", 1)
    add_paragraph(
        doc,
        "RiskRadar AI uses local sentence-transformer embeddings and a FAISS vector index to retrieve relevant policy context from four fraud policy documents. This keeps the investigation grounded in explicit policy evidence rather than relying only on model-generated reasoning.",
    )
    add_table(
        doc,
        ["Policy Document", "Role in Investigation"],
        [
            ["High_Value_Transaction_Policy.txt", "Guidance for large or unusual transaction amounts."],
            ["Velocity_and_Burst_Detection_Policy.txt", "Guidance for burst transaction patterns and rapid transaction frequency."],
            ["Geographic_Anomaly_Policy.txt", "Guidance for high-risk or unusual geographic signals."],
            ["Device_and_Payment_Instrument_Policy.txt", "Guidance for new device and instrument anomalies."],
        ],
        [3900, 5460],
    )
    add_image(
        doc,
        "docs/screenshots/Demo/07_sources_and_analyst_decision.jpg",
        "Figure 3. Policy sources and analyst decision controls.",
    )

    add_heading(doc, "8. Analyst Dashboard and Human Approval", 1)
    add_paragraph(
        doc,
        "The Streamlit dashboard is the primary analyst interface. It shows alert queue metrics, selected alert details, investigation results, risk reasoning, policy sources, agent trace, analyst decision controls, analyst note entry, and Decision History.",
    )
    for item in [
        "Alert queue removes handled transactions after decision save.",
        "Decision History shows recently handled alerts for audit and review.",
        "Analyst choices include Approve, Monitor, and Escalate.",
        "Analyst notes capture override reasoning or investigation context.",
    ]:
        add_bullet(doc, item)
    add_image(
        doc,
        "docs/screenshots/Demo/10_decision_history.jpg",
        "Figure 4. Decision History after analyst feedback is saved.",
    )

    add_heading(doc, "9. AWS Deployment and S3 Storage", 1)
    add_paragraph(
        doc,
        "The dashboard was deployed on AWS EC2 and exposed through Streamlit port 8501. S3-backed storage was added for the synthetic dataset, policy documents, and analyst feedback log. The EC2 instance uses an IAM role to access S3 without storing AWS access keys on the server.",
    )
    add_key_value_table(
        doc,
        [
            ("Compute", "AWS EC2, Ubuntu 24.04 LTS, t3.micro"),
            ("Application port", "8501"),
            ("Storage", "Amazon S3 bucket riskradar-ai-storage-alkarth with prefix riskradar/"),
            ("Security", "IAM role for S3 access; .env and .pem files ignored by git"),
            ("Runtime mode", "USE_S3_STORAGE=true and USE_LLM_SUMMARY=false for deterministic demos"),
        ],
    )
    add_image(
        doc,
        "docs/screenshots/aws_deployment/aws_03_live_app_home.png",
        "Figure 5. Live deployed RiskRadar AI dashboard on AWS.",
        width_in=5.6,
    )
    add_image(
        doc,
        "docs/screenshots/aws_deployment/aws_04_live_investigation_sources.png",
        "Figure 6. Live AWS investigation result with policy sources.",
        width_in=5.6,
    )

    add_heading(doc, "10. Optional Docker Containerization", 1)
    add_paragraph(
        doc,
        "Docker support was added as an optional reproducibility path. The project intentionally uses direct Docker CLI commands rather than Docker Compose, keeping the RFP containerization story simple while still allowing another user to build and run the app without a local Python setup.",
    )
    add_code_block(
        doc,
        [
            "docker build -t riskradar-ai .",
            'docker run --rm -p 8501:8501 --env-file .env -v \"$PWD/data:/app/data\" riskradar-ai',
        ],
    )
    add_paragraph(
        doc,
        "The data folder is mounted into the container so analyst feedback logs persist on the host machine. The .dockerignore file excludes secrets, keys, git metadata, caches, and local virtual environments from the image.",
    )

    add_heading(doc, "11. Testing and Verification", 1)
    add_table(
        doc,
        ["Verification Area", "Result"],
        [
            ["Static compile check", "Passed for app, graph, fraud, RAG, utils, tests, and root helper scripts."],
            ["Alert generation", "Passed and generated a working alert queue from synthetic transactions."],
            ["Risk engine", "Passed high, medium, and low risk scenario checks."],
            ["LangGraph workflow", "Passed script-style workflow validation."],
            ["End-to-end app flow", "Passed selected alert investigation and feedback logging checks."],
            ["AWS deployment", "Public app responded and live investigation displayed policy sources."],
            ["Docker", "Image build, container startup, and localhost Streamlit load were verified."],
        ],
        [3000, 6360],
    )
    add_paragraph(
        doc,
        "Current tests are script-style health checks rather than full pytest-collected test functions. This is documented as a known limitation and does not block the demo workflow.",
    )

    add_heading(doc, "12. RFP Alignment", 1)
    add_table(
        doc,
        ["Requirement Area", "Status", "Evidence"],
        [
            ["Fraud investigation assistant", "Complete", "Streamlit dashboard and workflow-driven investigation."],
            ["Agentic workflow", "Complete", "LangGraph specialist nodes and agent trace."],
            ["Policy-grounded evidence", "Complete", "FAISS retrieval from policy documents with source display."],
            ["Risk scoring", "Complete", "Weighted rule and policy scoring with risk classification."],
            ["Human approval", "Complete", "Analyst decision buttons, notes, feedback log, and Decision History."],
            ["Frontend/dashboard", "Complete", "Streamlit alert queue, investigation output, and decision workflow."],
            ["AWS deployment", "Complete", "EC2 public app, S3-backed data, policy, and feedback storage."],
            ["Optional containerization", "Complete", "Dockerfile, .dockerignore, README instructions, verified local run."],
            ["Final submission assets", "In progress", "Report, PPT, demo recording, and packaging remain."],
        ],
        [2500, 1500, 5360],
    )

    add_heading(doc, "13. Limitations and Future Scope", 1)
    for item in [
        "Feedback is currently stored as CSV/S3 file storage rather than a production database.",
        "Tests are script-style checks and can be converted into a fuller pytest suite.",
        "Historical case retrieval is not yet implemented.",
        "Investigation report export from the UI is not yet implemented.",
        "Workflow visualization is documented but not rendered interactively inside the dashboard.",
        "Security Group access should be tightened after demo use if the public app is no longer needed.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "14. Conclusion", 1)
    add_paragraph(
        doc,
        "RiskRadar AI demonstrates a complete AI-assisted fraud investigation workflow: alert triage, policy-grounded investigation, explainable recommendations, human approval, feedback capture, cloud deployment, and optional Docker reproducibility. The implementation is ready for final presentation packaging and demo recording, with the internal readiness target set for 25 June 2026 ahead of the updated 30 June 2026 deadline.",
    )

    add_heading(doc, "Appendix A. Key Repository Files", 1)
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ["streamlit_app/app.py", "Analyst dashboard and human decision workflow."],
            ["src/graph/workflow.py", "LangGraph workflow assembly."],
            ["src/graph/nodes.py", "Investigation node logic."],
            ["src/fraud/risk_engine.py", "Policy trigger and risk scoring logic."],
            ["src/rag/rag_pipeline.py", "FAISS policy retrieval pipeline."],
            ["src/utils/s3_storage.py", "S3 storage helper functions."],
            ["Dockerfile", "Optional Docker image definition."],
            ["README.md", "Setup, Streamlit, and Docker run instructions."],
        ],
        [3300, 6060],
    )

    add_heading(doc, "Appendix B. Submission Checklist", 1)
    for item in [
        "Source code repository.",
        "Synthetic transaction dataset.",
        "Fraud policy documents.",
        "Final report.",
        "Final PPT.",
        "Demo recording.",
        "AWS deployment proof screenshots.",
        "Dockerfile and Docker run instructions.",
    ]:
        add_bullet(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
